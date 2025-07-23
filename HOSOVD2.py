import tkinter as tk
from tkinter import ttk, messagebox
from docx import Document
import os
import unicodedata
from datetime import datetime
import re
from copy import deepcopy

# --- Bỏ dấu tiếng Việt và viết hoa, in liền ---
def bo_dau(text):
    return ''.join(c for c in unicodedata.normalize('NFKD', text) if not unicodedata.combining(c)).upper().replace(" ", "")

# --- Hàm thay thế cho file 2 và 4 (gộp text, giữ run đầu) ---
def replace_placeholders_simple(doc_path, replacements, output_path):
    doc = Document(doc_path)

    def replace_in_paragraph(paragraph, replacements):
        full_text = ''.join(run.text for run in paragraph.runs)
        new_text = full_text
        for k, v in replacements.items():
            new_text = new_text.replace(k, v)
        if new_text != full_text:
            for run in paragraph.runs:
                run.text = ''
            if paragraph.runs:
                paragraph.runs[0].text = new_text

    for p in doc.paragraphs:
        replace_in_paragraph(p, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p, replacements)

    doc.save(output_path)

# --- Hàm thay thế riêng cho file 3 giữ nguyên style run từng phần (giữ bố cục) ---
def replace_placeholders_file3(doc_path, replacements, output_path):
    doc = Document(doc_path)

    def copy_run_style(source_run, target_run):
        target_run.bold = source_run.bold
        target_run.italic = source_run.italic
        target_run.underline = source_run.underline
        try:
            target_run.font.color.rgb = source_run.font.color.rgb if source_run.font.color.rgb else None
        except Exception:
            pass
        try:
            target_run.font.name = source_run.font.name
        except Exception:
            pass
        try:
            target_run.font.size = source_run.font.size
        except Exception:
            pass

    def replace_in_paragraph(paragraph, replacements):
        # Nối toàn bộ text trong paragraph
        full_text = ''.join(run.text for run in paragraph.runs)
        # Thay thế toàn bộ placeholder
        for key, val in replacements.items():
            full_text = full_text.replace(key, val)
        # Xóa toàn bộ run cũ
        for run in paragraph.runs:
            run.text = ''
        n = len(paragraph.runs)
        if n == 0:
            paragraph.add_run(full_text)
            return
        # Chia đều text mới theo số run ban đầu để giữ style
        part_len = max(1, len(full_text) // n)
        parts = []
        start = 0
        for i in range(n - 1):
            parts.append(full_text[start:start+part_len])
            start += part_len
        parts.append(full_text[start:])
        # Gán lại từng phần text cho run tương ứng với style run gốc
        for i, run in enumerate(paragraph.runs):
            run.text = parts[i] if i < len(parts) else ''

    for p in doc.paragraphs:
        replace_in_paragraph(p, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p, replacements)

    doc.save(output_path)

# --- Hàm thay thế đơn giản cho file số 1 ---
def replace_placeholders_file1(doc_path, replacements, output_path):
    doc = Document(doc_path)

    def replace_in_paragraph_simple(paragraph, replacements):
        full_text = ''.join(run.text for run in paragraph.runs)
        replaced_text = full_text
        for key, val in replacements.items():
            if key in replaced_text:
                replaced_text = replaced_text.replace(key, val)
        if replaced_text != full_text:
            for run in paragraph.runs:
                run.text = ''
            if paragraph.runs:
                paragraph.runs[0].text = replaced_text
            else:
                paragraph.add_run(replaced_text)

    for p in doc.paragraphs:
        replace_in_paragraph_simple(p, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph_simple(p, replacements)

    doc.save(output_path)

# --- Tạo hồ sơ ---
def tao_ho_so():
    hoten = entry_hoten.get().strip()
    gioitinh = combo_gioitinh.get()
    diachi = entry_diachi.get().strip()
    soqd = entry_soqd.get().strip()
    cccd = entry_cccd.get().strip()
    cccd_d = entry_cccd_d.get().strip()
    cccd_m = entry_cccd_m.get().strip()
    cccd_y = entry_cccd_y.get().strip()
    qd_d = entry_qd_d.get().strip()
    qd_m = entry_qd_m.get().strip()
    qd_y = entry_qd_y.get().strip()

    fields = [hoten, gioitinh, diachi, soqd, cccd, cccd_d, cccd_m, cccd_y, qd_d, qd_m, qd_y]
    if any(not field for field in fields):
        messagebox.showerror("Lỗi", "Vui lòng điền đầy đủ tất cả các ô nhập.")
        return

    if any(char.isdigit() for char in hoten):
        messagebox.showerror("Lỗi", "Họ tên không được chứa số.")
        return
    if not soqd.isdigit():
        messagebox.showerror("Lỗi", "Số quyết định chỉ được chứa số.")
        return
    if not cccd.isdigit() or len(cccd) != 12:
        messagebox.showerror("Lỗi", "CCCD phải gồm đúng 12 chữ số.")
        return
    if not (cccd_d.isdigit() and 1 <= int(cccd_d) <= 31):
        messagebox.showerror("Lỗi", "Ngày cấp CCCD không hợp lệ.")
        return
    if not (cccd_m.isdigit() and 1 <= int(cccd_m) <= 12):
        messagebox.showerror("Lỗi", "Tháng cấp CCCD không hợp lệ.")
        return
    if not (cccd_y.isdigit() and 1900 <= int(cccd_y) <= datetime.now().year):
        messagebox.showerror("Lỗi", "Năm cấp CCCD không hợp lệ.")
        return
    if not (qd_d.isdigit() and 1 <= int(qd_d) <= 31):
        messagebox.showerror("Lỗi", "Ngày ra quyết định không hợp lệ.")
        return
    if not (qd_m.isdigit() and 1 <= int(qd_m) <= 12):
        messagebox.showerror("Lỗi", "Tháng ra quyết định không hợp lệ.")
        return
    if not (qd_y.isdigit() and 1900 <= int(qd_y) <= datetime.now().year):
        messagebox.showerror("Lỗi", "Năm ra quyết định không hợp lệ.")
        return

    ho_ten_folder = bo_dau(hoten)
    output_dir = f"D:/Finally-PhuocLong/{ho_ten_folder}"
    os.makedirs(output_dir, exist_ok=True)

    replacements = {
        "[HOTEN]": hoten,
        "[GIOITINH]": gioitinh,
        "[DIACHI]": diachi,
        "[SOQD]": soqd,
        "[CCCD]": cccd,
        "[CCCD_D]": cccd_d,
        "[CCCD_M]": cccd_m,
        "[CCCD_Y]": cccd_y,
        "[QD_D]": qd_d,
        "[QD_M]": qd_m,
        "[QD_Y]": qd_y,
    }

    base_path = "D:/Project-PhuocLong"
    files = [
        ("1 PDX thu ly_FullName.docx", f"1_PDX_thu_ly_{ho_ten_folder}.docx"),
        ("2 TB thu ly_FullName.docx", f"2_TB_thu_ly_{ho_ten_folder}.docx"),
        ("3 QDXM_FullName.docx", f"3_QDXM_{ho_ten_folder}.docx"),
        ("4 PhieuTrinhThongBaoThuLy.docx", f"4_PhieuTrinhThongBaoThuLy_{ho_ten_folder}.docx"),
    ]

    try:
        # File 1 - thay thế đơn giản
        src1 = os.path.join(base_path, files[0][0])
        dst1 = os.path.join(output_dir, files[0][1])
        replace_placeholders_file1(src1, replacements, dst1)

        # File 2,4 - thay thế nhanh, giữ run đầu
        src2 = os.path.join(base_path, files[1][0])
        dst2 = os.path.join(output_dir, files[1][1])
        replace_placeholders_simple(src2, replacements, dst2)

        src4 = os.path.join(base_path, files[3][0])
        dst4 = os.path.join(output_dir, files[3][1])
        replace_placeholders_simple(src4, replacements, dst4)

        # File 3 - thay thế giữ nguyên style run từng phần (giữ bố cục)
        src3 = os.path.join(base_path, files[2][0])
        dst3 = os.path.join(output_dir, files[2][1])
        replace_placeholders_file3(src3, replacements, dst3)

        messagebox.showinfo("Thành công", f"Hồ sơ đã tạo tại:\n{output_dir}")
    except Exception as e:
        messagebox.showerror("Lỗi", f"Không thể tạo hồ sơ:\n{e}")

# === Giao diện đẹp hơn, chỉnh sửa để không dính lẫn nhau ===

root = tk.Tk()
root.title("TẠO BỘ HỒ SƠ VÀNH ĐAI 2")
root.geometry("820x670")
root.configure(bg="#ecf0f1")

# Header
top_frame = tk.Frame(root, bg="#2980b9")
top_frame.pack(fill="x")
tk.Label(top_frame, text="HỆ THỐNG TẠO HỒ SƠ VÀNH ĐAI 2", font=("Segoe UI", 22, "bold"), bg="#2980b9", fg="white").pack(pady=25)

form_frame = tk.Frame(root, bg="#ffffff", padx=30, pady=20, bd=2, relief="groove")
form_frame.pack(padx=30, pady=15, fill="both", expand=True)

def add_row(label, widget, row):
    tk.Label(form_frame, text=label, font=("Segoe UI", 12), bg="#ffffff", fg="#2c3e50").grid(row=row, column=0, sticky="e", padx=15, pady=10)
    widget.grid(row=row, column=1, padx=15, pady=10, sticky="w")

def create_entry(width=35):
    e = tk.Entry(form_frame, font=("Segoe UI", 12), width=width, relief="flat", bg="#fdfdfd", 
                 highlightthickness=2, highlightcolor="#3498db", highlightbackground="#bdc3c7")
    e.bind("<FocusIn>", lambda e: e.widget.config(highlightcolor="#2980b9"))
    e.bind("<FocusOut>", lambda e: e.widget.config(highlightcolor="#bdc3c7"))
    return e

def create_small_entry(parent):
    e = tk.Entry(parent, font=("Segoe UI", 12), width=5, relief="flat", bg="#fdfdfd", 
                 highlightthickness=2, highlightcolor="#3498db", highlightbackground="#bdc3c7", justify="center")
    e.bind("<FocusIn>", lambda e: e.widget.config(highlightcolor="#2980b9"))
    e.bind("<FocusOut>", lambda e: e.widget.config(highlightcolor="#bdc3c7"))
    return e

# Tạo các trường nhập
entry_hoten = create_entry()
add_row("Họ tên:", entry_hoten, 0)

combo_gioitinh = ttk.Combobox(form_frame, values=["Ông", "Bà", "Ông (Bà)"], state="readonly", font=("Segoe UI", 12), width=12)
combo_gioitinh.current(0)
add_row("Giới tính:", combo_gioitinh, 1)


entry_diachi = create_entry()
add_row("Địa chỉ:", entry_diachi, 2)

entry_soqd = create_entry(22)
add_row("Số Quyết định:", entry_soqd, 3)

# --- Ngày ra quyết định (Row 4) ---
tk.Label(form_frame, text="Ngày ra quyết định:", font=("Segoe UI", 12), bg="#ffffff", fg="#2c3e50")\
    .grid(row=4, column=0, sticky="e", padx=15, pady=10)

frame_qd = tk.Frame(form_frame, bg="#ffffff")
frame_qd.grid(row=4, column=1, sticky="w")

entry_qd_d = create_small_entry(frame_qd)
entry_qd_d.pack(side="left", padx=(0, 8))
tk.Label(frame_qd, text="/", font=("Segoe UI", 14), bg="#ffffff", fg="#2c3e50").pack(side="left")
entry_qd_m = create_small_entry(frame_qd)
entry_qd_m.pack(side="left", padx=8)
tk.Label(frame_qd, text="/", font=("Segoe UI", 14), bg="#ffffff", fg="#2c3e50").pack(side="left")
entry_qd_y = create_small_entry(frame_qd)
entry_qd_y.pack(side="left", padx=(8, 0))

# --- Số CCCD (Row 5) ---
entry_cccd = create_entry(22)
add_row("Số CCCD:", entry_cccd, 5)

# --- Ngày cấp CCCD (Row 6) ---
tk.Label(form_frame, text="Ngày cấp CCCD:", font=("Segoe UI", 12), bg="#ffffff", fg="#2c3e50")\
    .grid(row=6, column=0, sticky="e", padx=15, pady=10)

frame_cccd = tk.Frame(form_frame, bg="#ffffff")
frame_cccd.grid(row=6, column=1, sticky="w")

entry_cccd_d = create_small_entry(frame_cccd)
entry_cccd_d.pack(side="left", padx=(0, 8))
tk.Label(frame_cccd, text="/", font=("Segoe UI", 14), bg="#ffffff", fg="#2c3e50").pack(side="left")
entry_cccd_m = create_small_entry(frame_cccd)
entry_cccd_m.pack(side="left", padx=8)
tk.Label(frame_cccd, text="/", font=("Segoe UI", 14), bg="#ffffff", fg="#2c3e50").pack(side="left")
entry_cccd_y = create_small_entry(frame_cccd)
entry_cccd_y.pack(side="left", padx=(8, 0))

# Nút Tạo hồ sơ
btn_create = tk.Button(root, text="TẠO HỒ SƠ", font=("Segoe UI", 16, "bold"), bg="#27ae60", fg="white",
                       activebackground="#2ecc71", activeforeground="white", relief="flat", padx=40, pady=15,
                       cursor="hand2", command=tao_ho_so)
btn_create.pack(pady=30)

root.mainloop()
