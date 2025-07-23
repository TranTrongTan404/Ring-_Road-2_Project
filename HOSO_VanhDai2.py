import tkinter as tk
from tkinter import ttk, messagebox
from docx import Document
import os
import unicodedata
from datetime import datetime
import re
from copy import deepcopy

# --- Bỏ dấu tiếng Việt và viết hoa, in liền ---
def bo_dau(text):
    return ''.join(c for c in unicodedata.normalize('NFKD', text) if not unicodedata.combining(c)).upper().replace(" ", "")

# --- Thay thế các trường, giữ nguyên style từng run ---
def replace_placeholders(doc_path, replacements, output_path):
    doc = Document(doc_path)

    def copy_run_style(source_run, target_run):
        target_run.bold = source_run.bold
        target_run.italic = source_run.italic
        target_run.underline = source_run.underline
        try:
            target_run.font.color.rgb = source_run.font.color.rgb if source_run.font.color.rgb else None
        except Exception:
            pass
        try:
            target_run.font.name = source_run.font.name
        except Exception:
            pass
        try:
            target_run.font.size = source_run.font.size
        except Exception:
            pass

    def replace_in_paragraph(paragraph, replacements):
        keys = sorted(replacements.keys(), key=len, reverse=True)
        pattern = '(' + '|'.join(re.escape(k) for k in keys) + ')'

        new_runs = []
        for run in paragraph.runs:
            parts = re.split(pattern, run.text)
            for part in parts:
                if part in replacements:
                    new_runs.append((replacements[part], deepcopy(run)))
                else:
                    new_runs.append((part, deepcopy(run)))

        if not new_runs:
            return

        for run in paragraph.runs:
            run.text = ''

        paragraph.runs[0].text = new_runs[0][0]
        copy_run_style(new_runs[0][1], paragraph.runs[0])

        for text, style_run in new_runs[1:]:
            r = paragraph.add_run(text)
            copy_run_style(style_run, r)

    for p in doc.paragraphs:
        replace_in_paragraph(p, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p, replacements)

    doc.save(output_path)

# --- Hàm thay thế đơn giản cho file số 1, tránh mất style ---
def replace_placeholders_file1(doc_path, replacements, output_path):
    doc = Document(doc_path)

    def replace_in_paragraph_simple(paragraph, replacements):
        full_text = ''.join(run.text for run in paragraph.runs)
        replaced_text = full_text
        for key, val in replacements.items():
            if key in replaced_text:
                replaced_text = replaced_text.replace(key, val)
        if replaced_text != full_text:
            # Xóa toàn bộ text cũ
            for run in paragraph.runs:
                run.text = ''
            # Gán lại cho run đầu nếu có, hoặc tạo run mới
            if paragraph.runs:
                paragraph.runs[0].text = replaced_text
            else:
                paragraph.add_run(replaced_text)

    for p in doc.paragraphs:
        replace_in_paragraph_simple(p, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph_simple(p, replacements)

    doc.save(output_path)

# --- Tạo hồ sơ ---
def tao_ho_so():
    hoten = entry_hoten.get().strip()
    gioitinh = combo_gioitinh.get()
    diachi = entry_diachi.get().strip()
    soqd = entry_soqd.get().strip()
    cccd = entry_cccd.get().strip()
    cccd_d = entry_cccd_d.get().strip()
    cccd_m = entry_cccd_m.get().strip()
    cccd_y = entry_cccd_y.get().strip()

    fields = [hoten, gioitinh, diachi, soqd, cccd, cccd_d, cccd_m, cccd_y]
    if any(not field for field in fields):
        messagebox.showerror("Lỗi", "Vui lòng điền đầy đủ tất cả các ô nhập.")
        return

    if any(char.isdigit() for char in hoten):
        messagebox.showerror("Lỗi", "Họ tên không được chứa số.")
        return
    if not soqd.isdigit():
        messagebox.showerror("Lỗi", "Số quyết định chỉ được chứa số.")
        return
    if not cccd.isdigit() or len(cccd) != 12:
        messagebox.showerror("Lỗi", "CCCD phải gồm đúng 12 chữ số.")
        return
    if not (cccd_d.isdigit() and 1 <= int(cccd_d) <= 31):
        messagebox.showerror("Lỗi", "Ngày cấp CCCD không hợp lệ.")
        return
    if not (cccd_m.isdigit() and 1 <= int(cccd_m) <= 12):
        messagebox.showerror("Lỗi", "Tháng cấp CCCD không hợp lệ.")
        return
    if not (cccd_y.isdigit() and 1900 <= int(cccd_y) <= datetime.now().year):
        messagebox.showerror("Lỗi", "Năm cấp CCCD không hợp lệ.")
        return

    ho_ten_folder = bo_dau(hoten)
    output_dir = f"D:/Finally-PhuocLong/{ho_ten_folder}"
    os.makedirs(output_dir, exist_ok=True)

    replacements = {
        "[HOTEN]": hoten,
        "[GIOITINH]": gioitinh,
        "[DIACHI]": diachi,
        "[SOQD]": soqd,
        "[CCCD]": cccd,
        "[CCCD_D]": cccd_d,
        "[CCCD_M]": cccd_m,
        "[CCCD_Y]": cccd_y,
    }

    base_path = "D:/Project-PhuocLong"
    files = [
        ("1 PDX thu ly_FullName.docx", f"1_PDX_thu_ly_{ho_ten_folder}.docx"),
        ("2 TB thu ly_FullName.docx", f"2_TB_thu_ly_{ho_ten_folder}.docx"),
        ("3 QDXM_FullName.docx", f"3_QDXM_{ho_ten_folder}.docx"),
        ("4 PhieuTrinhThongBaoThuLy.docx", f"4_PhieuTrinhThongBaoThuLy_{ho_ten_folder}.docx"),
    ]

    try:
        # Dùng hàm đơn giản cho file số 1
        src1 = os.path.join(base_path, files[0][0])
        dst1 = os.path.join(output_dir, files[0][1])
        replace_placeholders_file1(src1, replacements, dst1)

        # Dùng hàm giữ style cho các file còn lại
        for src, dst in files[1:]:
            src_path = os.path.join(base_path, src)
            dst_path = os.path.join(output_dir, dst)
            replace_placeholders(src_path, replacements, dst_path)

        messagebox.showinfo("Thành công", f"Hồ sơ đã tạo tại:\n{output_dir}")
    except Exception as e:
        messagebox.showerror("Lỗi", f"Không thể tạo hồ sơ:\n{e}")

# === GIAO DIỆN ===
root = tk.Tk()
root.title("TẠO BỘ HỒ SƠ VÀNH ĐAI 2")
root.geometry("780x540")
root.configure(bg="#fff9f0")

top_frame = tk.Frame(root, bg="#c0392b")
top_frame.pack(fill="x")
tk.Label(top_frame, text="HỆ THỐNG TẠO HỒ SƠ VÀNH ĐAI 2", font=("Arial", 18, "bold"), bg="#c0392b", fg="white").pack(pady=15)

form_frame = tk.Frame(root, bg="#fff9f0")
form_frame.pack(pady=20)

def add_row(label, widget, row):
    tk.Label(form_frame, text=label, font=("Segoe UI", 11), bg="#fff9f0").grid(row=row, column=0, sticky="e", padx=10, pady=5)
    widget.grid(row=row, column=1, padx=10, pady=5)

entry_hoten = tk.Entry(form_frame, font=("Segoe UI", 11), width=40)
add_row("Họ tên:", entry_hoten, 0)

combo_gioitinh = ttk.Combobox(form_frame, values=["Ông", "Bà"], state="readonly", font=("Segoe UI", 11), width=10)
combo_gioitinh.current(0)
add_row("Giới tính:", combo_gioitinh, 1)

entry_diachi = tk.Entry(form_frame, font=("Segoe UI", 11), width=40)
add_row("Địa chỉ:", entry_diachi, 2)

entry_soqd = tk.Entry(form_frame, font=("Segoe UI", 11), width=20)
add_row("Số Quyết định:", entry_soqd, 3)

entry_cccd = tk.Entry(form_frame, font=("Segoe UI", 11), width=20)
add_row("Số CCCD:", entry_cccd, 4)

tk.Label(form_frame, text="Ngày cấp CCCD:", font=("Segoe UI", 11), bg="#fff9f0").grid(row=5, column=0, sticky="e", padx=10, pady=5)
frame_ngay = tk.Frame(form_frame, bg="#fff9f0")
frame_ngay.grid(row=5, column=1, padx=10, pady=5, sticky="w")

entry_cccd_d = tk.Entry(frame_ngay, width=4, font=("Segoe UI", 11))
entry_cccd_d.pack(side="left", padx=(0, 5))
tk.Label(frame_ngay, text="/", bg="#fff9f0").pack(side="left")
entry_cccd_m = tk.Entry(frame_ngay, width=4, font=("Segoe UI", 11))
entry_cccd_m.pack(side="left", padx=(5, 5))
tk.Label(frame_ngay, text="/", bg="#fff9f0").pack(side="left")
entry_cccd_y = tk.Entry(frame_ngay, width=6, font=("Segoe UI", 11))
entry_cccd_y.pack(side="left", padx=(5, 0))

tk.Button(root, text="TẠO HỒ SƠ", command=tao_ho_so, bg="#e74c3c", fg="white", font=("Arial", 13, "bold"), padx=30, pady=12).pack(pady=30)

root.mainloop()
